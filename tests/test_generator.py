"""文档生成器单元测试."""

from pathlib import Path

from docx import Document as DocxDocument

from src.generator import DocumentGenerator
from src.models import ContentItem, DocumentRequest


def test_cover_spacing():
    """测试封面空行布局."""
    request = DocumentRequest(
        title="测试标题",
        subtitle="测试副标题",
        content=[],
    )

    output_path = "output/test_cover_spacing.docx"
    generator = DocumentGenerator()
    generator.generate(request, output_path)

    doc = DocxDocument(output_path)
    paras = [p for p in doc.paragraphs]

    title_idx = None
    subtitle_idx = None
    date_idx = None
    for i, p in enumerate(paras):
        if p.text == "测试标题":
            title_idx = i
        elif p.text == "测试副标题":
            subtitle_idx = i
        elif "2026年" in p.text or "2025年" in p.text:
            date_idx = i

    assert title_idx is not None
    assert subtitle_idx is not None
    assert date_idx is not None

    empty_before_title = sum(
        1 for i in range(title_idx)
        if paras[i].style.name == "Normal" and not paras[i].text.strip()
    )
    assert empty_before_title == 8

    empty_between_title_sub = sum(
        1 for i in range(title_idx + 1, subtitle_idx)
        if paras[i].style.name == "Normal" and not paras[i].text.strip()
    )
    assert empty_between_title_sub == 2

    empty_between_sub_date = sum(
        1 for i in range(subtitle_idx + 1, date_idx)
        if paras[i].style.name == "Normal" and not paras[i].text.strip()
    )
    assert empty_between_sub_date == 2

    print("✅ 封面布局测试通过")


def test_toc_inserted():
    """测试目录是否正确插入."""
    request = DocumentRequest(
        title="测试文档",
        content=[
            ContentItem(type="heading_1", text="一、第一章"),
            ContentItem(type="body_text", text="第一章内容。"),
            ContentItem(type="heading_2", text="（一）第一节"),
            ContentItem(type="body_text", text="第一节内容。"),
            ContentItem(type="heading_1", text="二、第二章"),
            ContentItem(type="body_text", text="第二章内容。"),
        ],
    )

    output_path = "output/test_toc.docx"
    generator = DocumentGenerator()
    generator.generate(request, output_path)

    doc = DocxDocument(output_path)

    # 找到目录标题
    toc_title_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text == "目录":
            toc_title_idx = i
            break
    assert toc_title_idx is not None, "未找到目录标题"
    assert doc.paragraphs[toc_title_idx].style.name == "目录标题"
    print(f"✅ 目录标题位置: [{toc_title_idx}] style={doc.paragraphs[toc_title_idx].style.name}")

    # 检查目录后是否有 TOC 域
    toc_para = doc.paragraphs[toc_title_idx + 1]
    has_toc_field = any(
        elem.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText'
        for run in toc_para.runs
        for elem in run._element
    )
    assert has_toc_field, "未找到 TOC 域"
    print("✅ TOC 域已插入")

    # 检查 updateFields 设置
    settings = doc.settings.element
    update_fields = settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields')
    assert update_fields is not None
    assert update_fields.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == 'true'
    print("✅ updateFields 已设置")

    print("✅ 目录测试通过")


def test_sections_and_page_numbers():
    """测试分节和页码：第一节无页码，第二节有页码."""
    request = DocumentRequest(
        title="分节测试",
        content=[
            ContentItem(type="heading_1", text="一、第一章"),
            ContentItem(type="body_text", text="第一章内容。"),
        ],
    )

    output_path = "output/test_sections.docx"
    generator = DocumentGenerator()
    generator.generate(request, output_path)

    doc = DocxDocument(output_path)

    # 检查节数
    assert len(doc.sections) == 2, f"应该有 2 节，实际有 {len(doc.sections)}"
    print(f"✅ 文档分节数: {len(doc.sections)}")

    # 第一节页脚应该为空
    s0_footer_text = doc.sections[0].footer.paragraphs[0].text
    s0_even_footer_text = doc.sections[0].even_page_footer.paragraphs[0].text
    assert s0_footer_text == "", f"第一节默认页脚应为空，实际为 {s0_footer_text!r}"
    assert s0_even_footer_text == "", f"第一节偶数页脚应为空，实际为 {s0_even_footer_text!r}"
    print("✅ 第一节页脚为空")

    # 第二节页脚应该有页码
    s1_footer_text = doc.sections[1].footer.paragraphs[0].text
    s1_even_footer_text = doc.sections[1].even_page_footer.paragraphs[0].text
    assert "—" in s1_footer_text, f"第二节默认页脚应包含页码，实际为 {s1_footer_text!r}"
    assert "—" in s1_even_footer_text, f"第二节偶数页脚应包含页码，实际为 {s1_even_footer_text!r}"
    print(f"✅ 第二节默认页脚: {s1_footer_text!r}")
    print(f"✅ 第二节偶数页脚: {s1_even_footer_text!r}")

    # 检查第二节是否有 pgNumType start=1
    sectPr = doc.sections[1]._sectPr
    pgNumType = sectPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgNumType')
    assert pgNumType is not None
    start = pgNumType.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}start')
    assert start == "1", f"起始页码应为 1，实际为 {start}"
    print("✅ 第二节起始页码为 1")

    # 检查奇偶页不同
    evenAndOddHeaders = sectPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}evenAndOddHeaders')
    assert evenAndOddHeaders is not None
    print("✅ 第二节启用奇偶页不同")

    print("✅ 分节和页码测试通过")


def test_generate_full_document():
    """测试生成一份包含所有元素类型的完整文档."""
    request = DocumentRequest(
        title="关于XXX项目的申报书",
        subtitle="XX单位",
        content=[
            ContentItem(type="heading_1", text="一、项目概述"),
            ContentItem(
                type="body_text",
                text="本项目旨在解决ABC-123领域的关键技术问题，计划于2025年底前完成。",
            ),
            ContentItem(type="heading_2", text="（一）研究背景"),
            ContentItem(
                type="body_text",
                text="随着信息技术的快速发展，数字化转型已成为各行各业的必然趋势。",
            ),
            ContentItem(type="heading_3", text="1. 国内现状"),
            ContentItem(type="body_text", text="国内已有多个城市开展了类似试点。"),
            ContentItem(type="heading_4", text="（1）北京试点"),
            ContentItem(type="body_text", text="北京市于2024年率先完成了系统建设。"),
            ContentItem(type="heading_5", text="1.1 子项目"),
            ContentItem(type="body_text", text="子项目描述。"),
            ContentItem(type="heading_6", text="1.1.1 细节"),
            ContentItem(type="body_text", text="细节描述。"),
            ContentItem(type="heading_7", text="1.1.1.1 更细"),
            ContentItem(type="body_text", text="更细的描述。"),
            ContentItem(type="heading_1", text="二、项目团队"),
            ContentItem(
                type="table",
                header=["姓名", "职务", "联系方式"],
                rows=[
                    ["张三", "项目负责人", "zhangsan@example.com"],
                    ["李四", "技术骨干", "lisi@example.com"],
                ],
            ),
            ContentItem(type="page_break"),
            ContentItem(type="heading_1", text="三、经费预算"),
            ContentItem(
                type="body_text",
                text="本项目总预算为500万元人民币，分三年拨付。",
            ),
        ],
    )

    output_path = "output/test_full.docx"
    generator = DocumentGenerator()
    result = generator.generate(request, output_path)

    assert Path(result).exists()
    print(f"✅ 文档生成成功: {result}")


def test_styles_applied():
    """测试各元素是否正确应用了模板样式."""
    request = DocumentRequest(
        title="测试标题",
        subtitle="测试副标题",
        content=[
            ContentItem(type="heading_1", text="一级标题"),
            ContentItem(type="heading_2", text="二级标题"),
            ContentItem(type="heading_3", text="三级标题"),
            ContentItem(type="heading_4", text="四级标题"),
            ContentItem(type="heading_5", text="五级标题"),
            ContentItem(type="heading_6", text="六级标题"),
            ContentItem(type="heading_7", text="七级标题"),
            ContentItem(type="body_text", text="正文内容"),
        ],
    )

    output_path = "output/test_styles.docx"
    generator = DocumentGenerator()
    generator.generate(request, output_path)

    doc = DocxDocument(output_path)
    body_paras = [p for p in doc.paragraphs if p.text.strip()]

    # 跳过封面的 3 段 + 目录的 2 段
    cover_count = 3
    toc_count = 2
    skip = cover_count + toc_count

    expected_styles = [
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
        "Heading 7",
        "Normal",
    ]

    for idx, expected in enumerate(expected_styles):
        actual = body_paras[skip + idx].style.name
        assert actual == expected, f"段落[{idx}] 样式不匹配: {actual} != {expected}"
        print(f"✅ [{idx}] {body_paras[skip + idx].text!r} -> style={actual}")


if __name__ == "__main__":
    test_cover_spacing()
    test_toc_inserted()
    test_sections_and_page_numbers()
    test_styles_applied()
    test_generate_full_document()
    print("\n✅ 所有测试通过!")
