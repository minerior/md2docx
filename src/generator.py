"""文档生成器核心：JSON -> docx（基于 Normal.docx 模板样式）."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.hdrftr import FooterPart
from docx.shared import Pt

from src.config import GongwenConfig
from src.models import ContentItem, DocumentRequest


class DocumentGenerator:
    """基于 Normal.docx 模板生成公文格式 Word 文档.

    核心逻辑：所有字体、字号、行距等格式由模板中的 Word 样式控制，
    代码只负责创建段落/表格并赋予正确的样式名。

    文档结构：
      第一节：封面 + 目录（无页码）
      第二节：正文（页码从 1 开始，奇右偶左）
    """

    def __init__(self, template_path: str = "Normal.docx"):
        self.config = GongwenConfig()
        self.template_path = Path(template_path)

    def generate(self, request: DocumentRequest, output_path: str) -> str:
        """生成文档并保存到 output_path，返回文件路径."""
        # 1. 加载模板
        if self.template_path.exists():
            doc = Document(str(self.template_path))
        else:
            doc = Document()

        # 2. 渲染封面（第一节）
        cover_cfg = self.config.get_cover_config()
        if cover_cfg.get("enabled", True):
            self._render_cover(doc, request)
            if cover_cfg.get("add_page_break", False):
                doc.add_page_break()

        # 3. 渲染目录（第一节）
        toc_cfg = self.config.get("toc", default={})
        if toc_cfg.get("enabled", False):
            self._render_toc(doc, toc_cfg)
            # 目录后的分页由分节符自动处理，不再需要单独的 page_break

        # 4. 添加分节符，开始第二节（正文）
        body_section = doc.add_section(start_type=WD_SECTION_START.NEW_PAGE)

        # 5. 清空第一节页脚（封面+目录不显示页码）
        self._clear_footer(doc.sections[0])

        # 6. 在第二节插入页码（从 1 开始，奇右偶左）
        self._insert_page_numbers(body_section)

        # 7. 渲染正文内容
        type_mapping = self.config.get_type_mapping()
        for item in request.content:
            style_name = type_mapping.get(item.type)
            if not style_name:
                continue

            if item.type == "page_break":
                doc.add_page_break()
            elif item.type == "table":
                self._render_table(doc, item)
            else:
                self._render_paragraph(doc, item, style_name)

        # 8. 保存
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return str(output_path)

    # ---------- 封面 ----------

    def _render_cover(self, doc: Document, request: DocumentRequest) -> None:
        cover_cfg = self.config.get_cover_config()
        spacing = cover_cfg.get("spacing", {})

        # 标题前空行
        for _ in range(spacing.get("before_title_lines", 0)):
            doc.add_paragraph(style="Normal")

        # 标题
        self._add_styled_paragraph(doc, request.title, cover_cfg.get("title_style", "Title"))

        # 副标题
        if request.subtitle:
            for _ in range(spacing.get("title_to_subtitle_lines", 0)):
                doc.add_paragraph(style="Normal")
            self._add_styled_paragraph(doc, request.subtitle, cover_cfg.get("subtitle_style", "Subtitle"))

        # 日期
        date_str = request.date or self._today_chinese()
        if date_str:
            gap_lines = spacing.get("subtitle_to_date_lines", 0) if request.subtitle else spacing.get("title_to_subtitle_lines", 0)
            for _ in range(gap_lines):
                doc.add_paragraph(style="Normal")
            self._add_styled_paragraph(doc, date_str, cover_cfg.get("date_style", "Subtitle"))

    @staticmethod
    def _today_chinese() -> str:
        """返回中文格式日期，如 2025年4月29日."""
        now = datetime.now()
        return f"{now.year}年{now.month}月{now.day}日"

    # ---------- 正文段落 ----------

    def _render_paragraph(self, doc: Document, item: ContentItem, style_name: str) -> None:
        self._add_styled_paragraph(doc, item.text, style_name)

    def _add_styled_paragraph(self, doc: Document, text: str, style_name: str) -> None:
        """创建段落并应用模板样式."""
        p = doc.add_paragraph()
        try:
            p.style = doc.styles[style_name]
        except KeyError:
            p.style = doc.styles["Normal"]
        if text:
            p.add_run(text)

    # ---------- 目录 ----------

    def _render_toc(self, doc: Document, toc_cfg: dict) -> None:
        """插入目录域."""
        # 目录标题
        title = toc_cfg.get("title", "目录")
        title_style = toc_cfg.get("title_style", "Heading 1")
        self._add_styled_paragraph(doc, title, title_style)

        # TOC 域
        levels = toc_cfg.get("levels", "1-7")
        paragraph = doc.add_paragraph()

        # begin fldChar
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        # instrText
        run = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' TOC \\o \"{levels}\" \\h \\z \\u '
        run._r.append(instrText)

        # separate fldChar
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)

        # 占位文本
        run = paragraph.add_run('（请右键此处选择"更新域"以生成目录）')

        # end fldChar
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

        # 设置 updateFields，让 Word 打开时提示更新域
        settings = doc.settings.element
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            settings.append(update_fields)
        update_fields.set(qn('w:val'), 'true')

    # ---------- 表格 ----------

    def _render_table(self, doc: Document, item: ContentItem) -> None:
        if not item.header or not item.rows:
            return

        num_cols = len(item.header)
        num_rows = len(item.rows) + 1  # +1 for header
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # 填充表头
        for col_idx, header_text in enumerate(item.header):
            cell = table.cell(0, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = doc.styles["表头"]  # 表头：黑体，无首行缩进
            p.add_run(header_text)

        # 填充数据行
        for row_idx, row_data in enumerate(item.rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= num_cols:
                    break
                cell = table.cell(row_idx + 1, col_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                p.style = doc.styles["No Spacing"]  # 内容：仿宋，无首行缩进
                p.add_run(cell_text)

        # 设置边框
        self._set_table_borders(table)

        # 撑满版心
        self._set_table_width(table, doc)

    @staticmethod
    def _set_table_borders(table) -> None:
        """设置表格标准全框线."""
        from docx.oxml import parse_xml

        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            '<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        tblBorders = parse_xml(
            r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>'
        )
        tblPr.append(tblBorders)

    @staticmethod
    def _set_table_width(table, doc: Document) -> None:
        """让表格宽度撑满版心可用宽度."""
        section = doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        available_width = page_width - left_margin - right_margin

        num_cols = len(table.columns)
        if num_cols == 0:
            return

        col_width = available_width / num_cols
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width

    # ---------- 页码 ----------

    def _clear_footer(self, section) -> None:
        """清空 section 的所有页脚内容."""
        for footer_obj in [section.footer, section.even_page_footer, section.first_page_footer]:
            for p in footer_obj.paragraphs:
                p.clear()

    def _insert_page_numbers(self, section) -> None:
        """在指定 section 中插入页码：奇偶页不同，单右双左，格式 — PAGE —."""
        # 1. 创建独立的 footer 部件
        default_footer_part = FooterPart.new(section.part.package)
        rId_default = section.part.relate_to(default_footer_part, RT.FOOTER)

        even_footer_part = FooterPart.new(section.part.package)
        rId_even = section.part.relate_to(even_footer_part, RT.FOOTER)

        # 2. 关联到 section
        sectPr = section._sectPr
        for ref_type, rId in [("default", rId_default), ("even", rId_even)]:
            footerRef = OxmlElement('w:footerReference')
            footerRef.set(qn('w:type'), ref_type)
            footerRef.set(qn('r:id'), rId)
            sectPr.append(footerRef)

        # 3. 启用奇偶页不同
        evenAndOddHeaders = sectPr.find(qn('w:evenAndOddHeaders'))
        if evenAndOddHeaders is None:
            evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
            sectPr.insert(0, evenAndOddHeaders)
        evenAndOddHeaders.set(qn('w:val'), 'true')

        # 4. 设置起始页码为 1
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        pgNumType.set(qn('w:start'), '1')

        # 5. 写入页码内容
        self._write_page_number_field(section.footer, "right")
        self._write_page_number_field(section.even_page_footer, "left")

    @staticmethod
    def _write_page_number_field(footer, alignment: str) -> None:
        """在 footer 段落中写入 — PAGE — 格式的页码."""
        align_map = {
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
        }

        p = footer.paragraphs[0]
        p.clear()
        p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.RIGHT)

        # —
        run = p.add_run('— ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

        # PAGE 域
        run = p.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        run = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)

        run = p.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)

        run = p.add_run('1')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

        run = p.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

        # —
        run = p.add_run(' —')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
